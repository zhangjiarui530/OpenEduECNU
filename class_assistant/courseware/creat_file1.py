import streamlit as st
import os
import time
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.oxml.ns import qn
from localApp import generate_text_from_model
from extract_high_freq_words import extract_high_freq_words_from_file

doc = Document()


# 添加word一级标题
def add_MainHeading(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.size = Pt(14)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# 添加word正文内容
def add_body(text):
    content = (text)
    p = doc.add_paragraph(content)
    p.runs[0].font.size = Pt(12)
    p.runs[0].font.name = '宋体'
    p.runs[0]._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


# 创建大纲word文档
def create_file(course_name, course_alltime, course_labtime, course_sub, course_type, file):
    st.warning('提交成功，教学大纲文档生成中，请稍后。。。。。。')

    # 创建一个新的Word文档

    st.warning('大纲课程概要部分生成中。。。。。。')

    # 添加word标题，并设置字体为宋体
    title = doc.add_heading(course_name + '课程' + '教学大纲', level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER  # 标题居中
    for run in title.runs:
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.bold = True  # 标题加粗

    prompt1 = "将" + course_name + "翻译成英语，只输出英文结果即可，不要其他任何符号和描述"
    course_nameE = generate_text_from_model(prompt1)

    prompt2 = "输出" + course_name + "的先修课程，2-3门重要的即可，只输出结果，不要任何符号和描述"
    course_pre = generate_text_from_model(prompt2)

    # 定义信息字段和对应的示例数据
    course_info = {
        '课程名称': course_name,
        '英文名称': course_nameE,
        '学时': course_alltime,
        '实验学时': course_labtime,
        '课程性质': course_type,
        '适用专业': course_sub,
        '先修课程': course_pre,
    }

    # 添加一个表格，表格的行数是字段数，列数是2
    table = doc.add_table(rows=len(course_info), cols=2)

    # 设置表格样式（可选）
    table.style = 'Table Grid'

    # 填充信息字段和数据
    for i, (field, value) in enumerate(course_info.items()):
        row_cells = table.rows[i].cells
        for idx, text in enumerate([field, value]):
            paragraph = row_cells[idx].paragraphs[0]
            run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
            run.text = text
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            if idx == 0:  # 字段名加粗
                run.bold = True
            run.font.size = Pt(12)

    st.success('大纲课程概要部分生成成功！')

    high_freq_words = extract_high_freq_words_from_file(file, 20)

    # 添加一个空行
    doc.add_paragraph()

    st.warning('大纲课程说明部分生成中。。。。。。')

    # 添加课程说明标题
    add_MainHeading('一、课程说明')

    prompt3 = "输出" + course_name + "这一门大学课程的课程说明，课程说明分为三段话，每段100字，课程类型为" + course_type + ",适用专业为" + course_sub + """先用一段话介绍课程基本内容和学生培养目标，再用一段话介绍课程教学设计概述，最后用一段话课程教学方法概述，
        只输出三段结果即可，不要任何符号和描述，不需要小标题，段与段之间不要空行，总计不超过300字，
        回答不要出现“第一段：”类似的描述，
        以下为另一门课程的课程说明，供参考:
        本课程讲解系统的基本架构、服务模式和开发原理，让学生对计算机系统的认识从单机、本地系统上升到云系统，重点培养学生的思维和能力。本课程理论和实践并重，让学生在应用和开发过程中，真正理解系统的魅力所在。
        在实践过程中，既锻炼学生搭建底层能力，又着重培养学生开发应用的能力。
        课程是一门较为前沿、相关知识和技术随着工业界的发展不断演进的课程，需要实时关注最新动态，并结合实际融入到教学过程中，最终使得学生在进入行业前就具备基本能力。
        """ + "本课程" + course_name + f"中频繁出现的关键词包括{high_freq_words}。"
    course_anal = generate_text_from_model(prompt3)

    # 添加课程说明内容
    add_body(course_anal)

    st.success('大纲课程说明部分生成成功！')

    # 添加一个空行
    doc.add_paragraph()

    st.warning('大纲课程目标部分生成中。。。。。。')

    # 添加课程目标标题
    add_MainHeading('二、课程目标')

    prompt4 = "输出" + course_name + "这一门大学课程的课程目标，分为四个点，按照下面的格式，目标1：了解什么内容，目标2：理解什么内容，目标3：掌握什么内容，目标4：运用什么内容，只分四段输出四个目标，不要任何符号和描述"
    course_tar = generate_text_from_model(prompt4)

    # 添加课程目标内容
    add_body(course_tar)

    st.success('大纲课程目标部分生成成功！')

    # 添加一个空行
    doc.add_paragraph()

    st.warning('大纲教学内容与学时安排部分生成中。。。。。。')

    # 添加教学内容标题
    add_MainHeading('三、教学内容与学时安排')

    prompt5 = "逐个章节设计" + course_name + "这一门大学课程的课程内容，课程类型为" + course_type + ",适用专业为" + course_sub + "注意每章的学时一般为2-4，总和需等于" + course_alltime + "," + """
        一般设计12个章节左右，每一章的设计都需你逐一列出
        章与章之间空一行,其他不空行,按照下面的格式输出,不要任何描述
        第XXX章：XXX
        学时：XXX
        内容：1、XXX；2、XXX；3、XXX
        要求学生：XXX
        """ + "本课程" + course_name + f"中频繁出现的关键词包括{high_freq_words}。"
    course_txt = generate_text_from_model(prompt5)

    # 添加教学内容
    add_body(course_txt)

    st.success('大纲教学内容与学时安排部分生成成功！')

    # 添加一个空行
    doc.add_paragraph()

    st.warning('大纲教学方法部分生成中。。。。。。')

    # 添加教学方法标题
    add_MainHeading('四、教学方法')

    prompt6 = "设计" + course_name + "这一门大学课程的教学方式，不要任何符号和描述，只输出答案，200字以内，可分段但不要分点，不需要小标题" + """
    以下为另一门课程课程的教学方法，供参考:
    在组织方式上，课前学生通过提前分发的课件，预习相关知识；课中在线下教室帮助学生梳理和巩固知识点，并且着重讲解重点和难点内容；课后，学生在前半学期完成一系列上机实践，后半学期运用课程技术技术完成开发项目。课程将组织1次理论测验和1次项目答辩，达到合格水平方能通过考核。
    学生通过预习和线下课堂学习、巩固理论知识，通过进行大量上机实践锻炼能力，磨炼分析问题、解决问题的动手实践能力，并且在训练过程中进一步加深对于理论知识的认知，最终达到对于课程 “知其然并且知其所以然”的学习效果，为日后从事相关工作打下坚实的基础。
    """
    course_teachway = generate_text_from_model(prompt6)

    # 添加教学方法内容
    add_body(course_teachway)

    st.success('大纲教学方法部分生成成功！')

    # 添加一个空行
    doc.add_paragraph()

    st.warning('大纲考核方式部分生成中。。。。。。')

    # 添加考核方式标题
    add_MainHeading('五、考核方式')

    prompt7 = "设计" + course_name + "这一门大学课程的考核方式，不要任何符号和描述，只输出答案，200字左右，可分段但不要分点，不需要小标题。" + """
    以下为另一门课程的考核方法，供参考:
    本课程采用闭卷形式进行考核，学生需在课程结束时完成一份闭卷考试。考试内容涵盖课程中所有知识点，包括理论概念、算法原理、应用实践等。考试时间为2小时。
    评价学生的方式包括平时作业、实验报告和课堂表现等。其中，平时作业包括课堂讲义、习题练习和作业批改等，占总成绩的30%；实验报告包括课程设计报告和实验报告等，占总成绩的20%；课堂表现包括课堂提问、讨论和演讲等，占总成绩的10%；期末考核占总成绩的40%。
    """
    course_exam = generate_text_from_model(prompt7)

    # 添加考核方式内容
    add_body("本课程旨在培养学生的理论素养、动手能力和创新思维。" + course_exam)

    st.success('大纲考核方式部分生成成功！')

    # 保存文档
    doc.save('result\\' + course_name + '教学大纲.docx')

    st.success('大纲教学大纲文档生成成功！')

    time.sleep(1)

    os.startfile('result')
    os.startfile('result\\' + course_name + '教学大纲.docx')
